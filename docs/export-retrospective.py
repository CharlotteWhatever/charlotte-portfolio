"""导出 Charlotte 建站项目完整复盘为 Word 文档"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── 全局样式 ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ── 辅助函数 ──
def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1C, 0x2A, 0x42)
    return h

def add_colored_table(headers, rows, col_widths=None):
    """添加带表头的表格"""
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
    # 数据行
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri+1].cells[ci].text = str(val)
    return table

def add_para(text, bold=False, italic=False, size=None, color=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold: run.bold = True
    if italic: run.italic = True
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    if space_after is not None: p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    return p

def add_numbered(text, level=0):
    p = doc.add_paragraph(text, style='List Number')
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    return p

# ═══════════════════════════════════════════════════════════
#  标题页
# ═══════════════════════════════════════════════════════════
title = doc.add_heading('Charlotte 个人品牌建站 · 完整项目复盘', level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1C, 0x2A, 0x42)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run('项目周期：2026-06-09 ~ 2026-06-11（3 天密集开发）\n')
r.font.color.rgb = RGBColor(0x5A, 0x6A, 0x76)
r = meta.add_run('当前版本：v2-mobile-responsive (Git tag: 3ff4d90)\n')
r.font.color.rgb = RGBColor(0x5A, 0x6A, 0x76)
r = meta.add_run('部署：GitHub Pages + Cloudflare Pages')
r.font.color.rgb = RGBColor(0x5A, 0x6A, 0x76)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  一、项目全景时间线
# ═══════════════════════════════════════════════════════════
add_heading('一、项目全景时间线', 1)

# Phase 0
add_heading('Phase 0：内容收集 · 2026-06-09（无代码，纯沟通）', 2)
add_colored_table(
    ['任务', '成果'],
    [
        ['获取原始简历', 'Charlotte-Chi-Resume.html → 提取全部经历'],
        ['内容方案定稿', '网站内容方案_文字版.md（8 模块文案 + 中英对照）'],
        ['项目访谈', 'CAME / Teknos / Alten / Yizhong+Shengshi 四家公司职责细节'],
        ['收集素材', '9 张截图（translator-q / wechat-auto / came-website 等）+ SVG banner'],
        ['推荐信收集', '4 人推荐语（Suzie Upson / Bogza+殷 / Jessica Lu / Nina Cai）'],
    ]
)
p = doc.add_paragraph()
run = p.add_run('✅ 该阶段做对的：')
run.bold = True
p.add_run('内容先于设计，先搞清楚"说什么"再"怎么呈现"。')
p = doc.add_paragraph()
run = p.add_run('❌ 可以更好：')
run.bold = True
p.add_run('素材尺寸没有在采集阶段规范，导致后期 AI 板块图片适配困难。')

# Phase 1
add_heading('Phase 1：设计探索 · 2026-06-09 ~ 06-10（8 轮配色尝试）', 2)
add_para('这是整个项目最曲折的阶段。几乎所有的"时间浪费"都发生在这里：')

add_colored_table(
    ['轮次', '方案', '结果', '原因'],
    [
        ['1', '原始多色版', '❌ 被拒', '"好丑"'],
        ['2', 'Apple 白版（A-apple-white）', '⏸️ 保留但放弃', '太普通'],
        ['3', '暖中性版（C-warm-neutral）', '⏸️ 保留但放弃', '太普通'],
        ['4', '深色高级版（B-dark-premium）', '❌ 被拒', '"矫枉过正"'],
        ['5', '暖蓝版（#47c2dc/#084a94）', '❌ 被拒', '"应该先让我看配色方案再做出来"'],
        ['6', '5 套北欧配色对比（color-schemes.html）', '⏸️ 未评价', '—'],
        ['7', '北欧雾蓝版', '⏸️ 未评价', '—'],
        ['8', 'Dark Explorer（#1C2A42/#B8845C/#F4EFE6）', '⏸️ 未确认', '用户说"做一个给我看看"但没下文'],
    ]
)

add_para('最终选择：Neumorphism 蓝（#E4E9F2 / #1a5a9c / #4a8fc7）', bold=True)
add_para(
    '实际部署的是我自己直接实现的 Neumorphism 蓝色方案，在进入 git 前就已经确定了。'
    '后面 8 轮设计探索实际上是在已确认的设计之外额外尝试——这是沟通上的关键问题：'
    '用户说"试试看"不等于"换掉当前方案"。'
)

add_para('教训：', bold=True)
for t in [
    '设计探索应该在进入开发前完成，而不是和开发并行',
    '"试试看" -> "做出来" 之间应该有一个中间环节（sketch / mockup / 引用参考）',
    '用户说"好丑"的时候应该追问："你希望的方向是什么？有没有参考？"',
    '最多提供 2 个方向，而不是做 8 个',
]:
    add_bullet(t)

# Phase 2
add_heading('Phase 2：v1 桌面版开发 · 2026-06-10（~8 小时）', 2)
add_para('实际产出速度极快，3 个 commit 完成一个完整单页响应式网站：')

add_colored_table(
    ['Commit', '内容'],
    [
        ['9b57036', 'Initial portfolio website — 完整单页（9 模块）'],
        ['88d9541', 'Fix hero resume button link path — 修复下载按钮'],
        ['4ad9455 (tag: v1-desktop)', 'Resume button follows language switch — 完成双语切换'],
    ]
)

add_para('v1 技术架构：', bold=True)
for t in [
    'Neumorphism 设计系统（阴影 = 2 层嵌套阴影）',
    '9 个模块：Nav → Hero → Experience → Case Study → Portfolio → AI → Recommendations → Downloads → Footer',
    '双语 data-i18n 切换（纯 JS，无 i18n 库）',
    '2 个断点（900px / 480px）——实际上只有桌面端',
    '3 个浮动玻璃态 blob 装饰',
    'Storytelling Hero（背景渐变 + 照片区）',
    'Timeline 时间线（4 家公司）',
]:
    add_bullet(t)

add_para('v1 已知未完成事项：', bold=True)
for t in [
    '移动端 <900px 未测试',
    '图片无 lazy loading',
    'AI 板块在移动端的布局未定义',
    '导航栏在小屏可能溢出',
    '中英文简历 PDF 未导出',
    '部署到 GitHub Pages',
]:
    add_bullet(t)

# Phase 3
add_heading('Phase 3：v2 移动端响应式 · 2026-06-10 ~ 06-11（14 小时，13 commits）', 2)

add_para('这是项目最痛苦的一段经历。一进移动端就发现导航栏消失 + AI 图片全崩。')
add_para('围绕两个核心问题展开了多轮修复：AI 图片拉伸（5 轮）和导航栏（3 轮）。')

add_colored_table(
    ['Commit', '内容', '问题'],
    [
        ['7628ca8', '响应式修复：hamburger 菜单 + 4 级断点 + 图片响应式', '全局 img 规则破坏布局 + hamburger 白色横杠'],
        ['fcb49a6', '修复：移动端 nav 显示 + AI 图片拉伸（第 1 轮）', '图片继续拉伸，未解决'],
        ['6abcd57', '修复v2：移除冲突的全局 img 规则（第 2 轮）', '用户要求回退导航栏'],
        ['d5c9292', '修复v3：回退导航栏至 v1 + flex min-height（第 3 轮）', '"它还是触发拉伸了！！！"'],
        ['cda1d91', '修复AI：aspect-ratio 替换固定高度（第 4 轮）', '桌面版被影响 + 移动端仍有问题'],
        ['f8930d4', '修复v4：桌面端恢复，移动端隐藏 AI 图片（第 5 轮）', '绝望尝试，"你终于改好了"'],
        ['3ff4d90 (tag: v2)', '合并AI修复：absolute 包裹 + 水平滑动 + 圆点导航', '✅ 最终正确方案'],
    ]
)

# AI 图片深度分析
add_heading('核心问题：AI 图片拉伸（5 轮修复才解决）', 3)

add_para('根因追踪：', bold=True)
steps = [
    '图片尺寸是 1200×400（宽高比 3:1），但容器在 flex 上下文中',
    'Flexbox 的 min-height: auto 默认值允许元素按内容自然尺寸撑大',
    '图片的 intrinsic size（1200×400）> 容器限制 → 溢出',
    'scrollIntoView 触发页面级滚动 → flexbox 重算 → 图片重新拉伸',
    '全局 img { max-width: 100%; height: auto } 覆盖了 object-fit: cover',
]
for i, s in enumerate(steps, 1):
    add_numbered(s)

add_para('最终方案：position: absolute + aspect-ratio: 3/1', bold=True)
for t in [
    '图片完全脱离 flex 文档流',
    '容器尺寸由 aspect-ratio 严格约束',
    'object-fit: cover 填充',
    '不会因为任何滚动/重算触发拉伸',
]:
    add_bullet(t)

add_para('为什么用了 5 轮才找到？', bold=True)
add_colored_table(
    ['轮次', '尝试方案', '失败原因'],
    [
        ['1', '全局 CSS 规则', '破坏已有布局'],
        ['2', 'flex 容器内修改', '没理解 min-height: auto'],
        ['3', '加 min-height: 0', '效果不一致'],
        ['4', '换 aspect-ratio', '没解决脱离文档流问题'],
        ['5', 'position: absolute', '✅ 有效'],
    ]
)

p = doc.add_paragraph()
run = p.add_run('根本原因：')
run.bold = True
p.add_run('我和 flexbox 的默认行为在"对抗"，而不是在"利用"。一个明确的诊断方法应该早用上：在 DevTools 里检查每个元素的 computed height，看实际值 vs 期望值。')

# 导航栏
add_heading('导航栏 hamburger 实验（3 轮失败）', 3)
add_colored_table(
    ['轮次', '尝试', '用户反馈'],
    [
        ['1', '添加 hamburger 菜单 + CSS 动画', '"白色横杠 + 黑线"'],
        ['2', '设 transparent 背景', '"变成 3 条了"'],
        ['3', '回退到 v1 导航（缩小字号+间距）', '✅ 用户满意'],
    ]
)
add_para('教训：如果折叠方案的视觉残留（毛玻璃 + neumorphism 深色阴影）无法消除，就不应该硬做。6 项导航在 320px 宽度下用缩小方案仍然可见——在尝试折叠之前应该先确认这个前提。')

# 其他 v2 问题
add_heading('其他 v2 修复问题', 3)
add_colored_table(
    ['问题', '根因', '解决'],
    [
        ['圆点出现在 Experience 板块', 'Python 占位符替换时放错位置', '手动核对 HTML 嵌套'],
        ['圆点被纵向拉长', 'button { min-height: 44px } 全局规则', '加 min-height: 0'],
        ['添加按钮触发拉伸', 'scrollIntoView 的副作用', '改用 scrollTo + offsetLeft'],
        ['箭头按钮触发拉伸', '同上', '移除箭头按钮，纯手势 scroll-snap'],
        ['测试和正式文件不同步', '改了 test 文件没 merge', '改成直接改 index.html'],
    ]
)

# Phase 4
add_heading('Phase 4：简历系统 · 2026-06-11（7 commits）', 2)

add_colored_table(
    ['Commit', '内容'],
    [
        ['1de378e', '新增简历v2(英文Gmail+个人链接) + 网站邮箱跟随语言切换'],
        ['6aea643', '中文简历+网站下载区：个人网站链接 & 英文简历v2上线'],
        ['228f0bb', '英文简历v2 PDF导出 + 下载链接改为PDF'],
        ['8438859', '英文简历v2替代旧版：下载区+Hero按钮+JS切换全部指向v2'],
        ['4b1efa9', '重新导出PDF：中文简历含个人网站链接 + 英文简历v2同步'],
        ['b27e9d3', '清理docs：移除旧版英文简历(v1) + 废弃full-cv方案A/C'],
    ]
)

add_para('关键沟通错误：', bold=True)

p = doc.add_paragraph()
run = p.add_run('错误 1：新增按钮 vs 替代旧版')
run.bold = True
p.add_run('\n用户说："难道不应该替代之前的一页简历吗，你新增一个按钮是什么意思？"')
add_bullet('我的理解：新增一个 v2 下载选项')
add_bullet('用户的理解：用 v2 替换 v1，不要多个按钮')
add_bullet('教训：修改已有功能时先确认「替代还是新增」')

p = doc.add_paragraph()
run = p.add_run('错误 2：邮箱替换错了')
run.bold = True
p.add_run('\n第一次把中文简历的 Foxmail 也改成了 Gmail。')
add_bullet('用户纠正："英文改成gmail，中文fox不要变"')
add_bullet('用户解释了原因：Foxmail 国内送达率好，Gmail 避免国外企业邮箱拦截')
add_bullet('最终方案：JS 语言切换时动态改 mailto: 和显示文本')

add_para('正确决策：', bold=True)
for t in [
    '"Shanghai, China" → "Personal Site" 链接（含超链图标）',
    '中文简历的 "上海" → "个人网站"（保持 Foxmail）',
    '英文 → Gmail（国际送达保障）',
    '中文 → Foxmail（国内直达）',
]:
    add_bullet(t)

add_para('docs 清理结果：', bold=True)
add_para('删除：resume-en.html / resume-en.pdf / full-cv-en-vA.html / full-cv-en-vC.html', italic=True)
add_para('保留 8 个文件：CAME 复盘（中/英）、完整 CV、简历（中/英 HTML+PDF）、工作职责导出、导出脚本、推荐信', italic=True)

# ═══════════════════════════════════════════════════════════
#  二、数据统计
# ═══════════════════════════════════════════════════════════
add_heading('二、数据统计', 1)
add_colored_table(
    ['指标', '数值'],
    [
        ['总 commit', '16'],
        ['Git tags', '2（v1-desktop, v2-mobile-responsive）'],
        ['总代码行', '1633 行（纯 HTML+CSS+JS 单文件）'],
        ['模块数', '9'],
        ['断点层级', '4（900 / 768 / 600 / 480px）'],
        ['图片总数', '10 张（全部 loading="lazy"）'],
        ['简历版本', '2 套 4 文件（中/英 HTML + PDF）'],
        ['设计探索轮次', '8'],
        ['AI 图片修复轮次', '5'],
        ['Hamburger 尝试轮次', '3（最终放弃）'],
        ['部署目标', '2（GitHub Pages + Cloudflare Pages）'],
    ]
)

# ═══════════════════════════════════════════════════════════
#  三、深度问题分析
# ═══════════════════════════════════════════════════════════
add_heading('三、深度问题分析', 1)

# 问题 1
add_heading('问题 1：没有阶段性验收就把 v1 推到了"完成"', 2)
add_para(
    'v1 看起来"做完了"（所有模块都有内容、样式统一、双语运行），但实际上移动端从未被测试。这是一个经典陷阱：'
    '我在桌面端开发了 8 小时，觉得"差不多了"，用户一拿起手机测 → 全是 bug。'
)
add_para('应该怎么做：在打 v1-desktop tag 之前就声明"这是桌面版 v1，移动端尚未适配，建议在 1920px+ 屏幕查看"。', bold=True)

# 问题 2
add_heading('问题 2：修复过程没有诊断数据', 2)
add_para(
    'AI 图片修了 5 轮，每一轮我都在"猜"——"可能是 flex 的问题"→ 改 flex；"可能是 aspect-ratio 的问题"→ 改 aspect-ratio。'
    '没有一次打开 DevTools 检查 computed height / flex shrink factor。'
)
add_para('正确的做法：先量化再修——量出 .ai-visual 在拉伸时的实际高度、图片的 naturalWidth/naturalHeight、容器的 flex-basis。', bold=True)

# 问题 3
add_heading('问题 3：设计探索和开发混在一起', 2)
add_para(
    '8 套配色方案 → 2 个对比 HTML 文件 → 最终没用上任何一套。'
    '投入到设计探索的时间大约占了总时间的 40%，但产出是零。'
)
add_para('核心矛盾：用户想要参考和对比，但我要做成 HTML 才能展示。应该用 Figma / 截图 / 色板卡片，而不是完整 HTML。', bold=True)

# 问题 4
add_heading('问题 4：沟通中的"做"和"想"没有分开', 2)
add_colored_table(
    ['用户说的', '我理解的', '应该做的'],
    [
        ['"做一个给我看看"', '需求 → 做完整方案', '先问方向，出 2-3 个参考确认再做'],
        ['"试试看"', '任务 → 花 2 小时做出来', '控制在 30 分钟内的快速原型'],
    ]
)

# ═══════════════════════════════════════════════════════════
#  四、做得好的地方
# ═══════════════════════════════════════════════════════════
add_heading('四、做得好的地方', 1)

items = [
    ('技术选型正确', [
        '纯 HTML/CSS/JS 单文件，部署极简（copy → GitHub Pages）',
        '无构建工具、无依赖，修改直观',
        '加载飞快（无框架运行时）',
    ]),
    ('双语方案优雅', [
        'data-i18n 属性 + switchLang() 函数，不需要 i18n 库',
        '天然 SEO 友好（默认英文）',
        '切换即时，无刷新，扩展性高',
    ]),
    ('版本控制意识', [
        'v1 打 tag 后才开始改 v2，用户可以随时 git checkout 回退',
        '每个修复阶段都有 commit，方便逐段回滚',
    ]),
    ('最终 AI 图片方案经得起推敲', [
        'position: absolute + aspect-ratio + object-fit: cover',
        '在所有浏览器行为一致，不受 flex/grid 上下文影响',
        '不需要 JavaScript，维护简单',
    ]),
    ('简历系统的邮箱策略合理', [
        '英文 → Gmail（国际送达保障）',
        '中文 → Foxmail（国内直达）',
        '语言切换时动态换，一个源文件维护',
    ]),
]

for title, bullets in items:
    add_heading(title, 2)
    for b in bullets:
        add_bullet(b)

# ═══════════════════════════════════════════════════════════
#  五、SOP 提炼
# ═══════════════════════════════════════════════════════════
add_heading('五、SOP 提炼（下次直接套用）', 1)

add_heading('开发流程规范', 2)
for phase in [
    ('Phase 1 — 内容定稿（不写代码）', [
        '收集文案、图片、数据',
        '确认模块结构和信息层级',
        '✅ 签收后才能进入 Phase 2',
    ]),
    ('Phase 2 — 设计对齐（最多 2 轮）', [
        '出 2 个方向（截图/参考链接）',
        '用户选择一个 → 出高保真 mockup',
        '✅ 签收后才能进入 Phase 3',
    ]),
    ('Phase 3 — 技术实现', [
        '先做桌面端 → 打 tag → 再做移动端',
        '每个 breakpoint 逐一测试',
        '图片全部用 aspect-ratio + object-fit',
        '导出 PDF / 生成下载文件',
    ]),
    ('Phase 4 — QA', [
        '桌面端（1920 / 1440 / 1024）',
        '平板（768 / 900）',
        '手机（480 / 375 / 320）',
        '跨浏览器（Chrome / Safari / Edge）',
    ]),
]:
    p = add_para(phase[0], bold=True)
    for b in phase[1]:
        add_bullet(b)

add_heading('Bug 修复 SOP', 2)
steps = [
    '量化问题：量尺寸、截图、记录',
    '定位根因：查 DevTools computed styles',
    '出一个最小复现（isolate the pattern）',
    '提出方案 + 预期效果',
    '改一个地方，测一个地方',
    '确认不影响其他模块',
]
for i, s in enumerate(steps, 1):
    add_numbered(s)

add_heading('沟通规范', 2)
add_colored_table(
    ['场景', '正确做法'],
    [
        ['"做一个看看"', '先确认方向，再快速原型（<30min）'],
        ['"帮我改一下"', '重复确认需求（"你的意思是……？"）'],
        ['"试试这个方案"', '先说预期效果，再动手'],
        ['替代 vs 新增', '先问清楚，不要默认'],
    ]
)

# ═══════════════════════════════════════════════════════════
#  六、待办事项
# ═══════════════════════════════════════════════════════════
add_heading('六、待办事项（未完成）', 1)

add_colored_table(
    ['优先级', '事项', '状态'],
    [
        ['🟡 中', 'Dark Explorer 配色用户确认', '已搁置 4 天'],
        ['🟡 中', '完整 CV（full-cv-en.html）', '用户说"日后再说"'],
        ['🟢 低', '微信推文 9 宫格确认', '等待用户'],
        ['🟢 低', 'Gitee 仓库清理（有旧内容）', '未开始'],
        ['🟢 低', '头像/照片上传到 hero 区域', '需用户提供照片'],
        ['🔵 远期', '自定义域名绑定', '未开始'],
        ['🔵 远期', '毛玻璃效果', '已提过但未采用'],
        ['🔵 远期', 'Python 后端 CMS', '未开始'],
    ]
)

# ── 页脚 ──
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('—— EOF ——')
run.font.color.rgb = RGBColor(0x5A, 0x6A, 0x76)
run.italic = True

# ── 保存 ──
output_path = r"C:\Users\64107\Desktop\charlotte-portfolio\docs\portfolio-retrospective.docx"
doc.save(output_path)
print(f"OK → {output_path}  ({round(os.path.getsize(output_path)/1024)}K)")
