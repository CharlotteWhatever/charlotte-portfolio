"""导出 Charlotte 建站项目完整复盘 v2 — 含用户补充的 5 大块分析与 SOP"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── 全局样式 ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ── 辅助函数 ──
def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1C, 0x2A, 0x42)
    return h

def add_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri+1].cells[ci].text = str(val)
    return table

def add_para(text, bold=False, italic=False, size=None, color=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold: run.bold = True
    if italic: run.italic = True
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    if space_after is not None: p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    return p

def add_numbered(text, level=0):
    p = doc.add_paragraph(text, style='List Number')
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    return p

# ═══════════════════════════════════════════════════════════
#  标题页
# ═══════════════════════════════════════════════════════════
add_heading('Charlotte 个人品牌建站 · 完整项目复盘', 0)
for run in doc.paragraphs[0].runs:
    run.font.color.rgb = RGBColor(0x1C, 0x2A, 0x42)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run('版本：v2（含用户补充的 5 大块分析与系统化 SOP 提炼）\n')
r.font.color.rgb = RGBColor(0x5A, 0x6A, 0x76)
r = meta.add_run('项目周期：2026-06-09 ~ 2026-06-11（3 天密集开发）\n')
r.font.color.rgb = RGBColor(0x5A, 0x6A, 0x76)
r = meta.add_run('当前版本：v2-mobile-responsive (Git tag: 3ff4d90)  ·  部署：GitHub Pages + Cloudflare Pages')
r.font.color.rgb = RGBColor(0x5A, 0x6A, 0x76)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  一、项目全景时间线
# ═══════════════════════════════════════════════════════════
add_heading('一、项目全景时间线', 1)

add_heading('Phase 0：内容收集 · 2026-06-09（无代码，纯沟通）', 2)
add_table(
    ['任务', '成果'],
    [
        ['获取原始简历', 'Charlotte-Chi-Resume.html → 提取全部经历'],
        ['内容方案定稿', '网站内容方案_文字版.md（8 模块文案 + 中英对照）'],
        ['项目访谈', 'CAME / Teknos / Alten / Yizhong+Shengshi 四家公司职责细节'],
        ['收集素材', '9 张截图（translator-q / wechat-auto / came-website 等）+ SVG banner'],
        ['推荐信收集', '4 人推荐语（Suzie Upson / Bogza+殷 / Jessica Lu / Nina Cai）'],
    ]
)
p = doc.add_paragraph()
run = p.add_run('✅ 做对的：'); run.bold = True
p.add_run('内容先于设计，先搞清楚"说什么"再"怎么呈现"。')
p = doc.add_paragraph()
run = p.add_run('❌ 可以更好：'); run.bold = True
p.add_run('素材尺寸没有在采集阶段规范，导致后期 AI 板块图片适配困难。')

add_heading('Phase 1：设计探索 · 2026-06-09 ~ 06-10（8 轮配色尝试）', 2)
add_para('这是整个项目最曲折的阶段。几乎所有的时间浪费都发生在这里。')
add_table(
    ['轮次', '方案', '结果', '原因'],
    [
        ['1', '原始多色版', '❌ 被拒', '"好丑"'],
        ['2', 'Apple 白版（A-apple-white）', '⏸️ 保留但放弃', '太普通'],
        ['3', '暖中性版（C-warm-neutral）', '⏸️ 保留但放弃', '太普通'],
        ['4', '深色高级版（B-dark-premium）', '❌ 被拒', '"矫枉过正"'],
        ['5', '暖蓝版（#47c2dc/#084a94）', '❌ 被拒', '"应该先让我看配色方案再做出来"'],
        ['6', '5 套北欧配色对比', '⏸️ 未评价', '—'],
        ['7', '北欧雾蓝版', '⏸️ 未评价', '—'],
        ['8', 'Dark Explorer', '⏸️ 未确认', '—'],
    ]
)

add_para('⚠️ 重要事实补充：最终 Neumorphism 蓝方案能够成功，是因为用户在网上下载了一个固定的提示词模板，发给 AI 后按照那个模板做出来的。这不是 AI 独立的设计能力，而是用户提供了关键的设计参考输入。这一点在初版复盘中被遗漏了。', italic=True)

add_para('本轮教训：', bold=True)
for t in [
    '设计探索必须在开发前完成，不可并行',
    '"试试看" → "做出来"之间应有一个中间环节（截图/参考链接/色板卡片），而不是直接出完整 HTML',
    '用户说"好丑"时应追问方向，而不是继续出更多方案',
    '最多提供 2 个方向对比，而不是 8 个',
]:
    add_bullet(t)

add_heading('Phase 2：v1 桌面版开发 · 2026-06-10（~8 小时）', 2)
add_table(
    ['Commit', '内容'],
    [
        ['9b57036', 'Initial portfolio website — 完整单页（9 模块）'],
        ['88d9541', 'Fix hero resume button link path'],
        ['4ad9455 (tag: v1-desktop)', 'Resume button follows language switch — 完成双语切换'],
    ]
)

add_para('v1 架构：', bold=True)
for t in [
    'Neumorphism 设计系统（2 层嵌套阴影）',
    '9 个模块：Nav → Hero → Experience → Case Study → Portfolio → AI → Recommendations → Downloads → Footer',
    '双语 data-i18n 切换（纯 JS，无 i18n 库）',
    '2 个断点（900px / 480px）——实际上只适配了桌面端',
    '3 个浮动玻璃态 blob 装饰',
]:
    add_bullet(t)

add_para('v1 已知未完成事项：', bold=True)
for t in [
    '移动端 <900px 从未测试',
    '图片无 lazy loading',
    'AI 板块在移动端的布局未定义',
    '中英文简历 PDF 未导出',
    '未部署',
]:
    add_bullet(t)

add_heading('Phase 3：v2 移动端响应式 · 2026-06-10 ~ 06-11（14 小时，13 commits）', 2)
add_para('这是项目最痛苦的一段。一进移动端就发现导航栏消失 + AI 图片全崩。围绕两个核心问题展开了多轮修复。')

add_table(
    ['Commit', '内容', '问题'],
    [
        ['7628ca8', '响应式修复：hamburger + 4 级断点 + 图片响应式', '全局 img 规则破坏布局 + hamburger 白色横杠'],
        ['fcb49a6', '修复：移动端 nav + AI 图片（第 1 轮）', '图片继续拉伸'],
        ['6abcd57', '修复v2：移除全局 img 规则（第 2 轮）', '用户要求回退导航栏'],
        ['d5c9292', '修复v3：回退导航至 v1 + flex min-height（第 3 轮）', '"它还是触发拉伸了！！！"'],
        ['cda1d91', '修复AI：aspect-ratio（第 4 轮）', '桌面版被影响'],
        ['f8930d4', '修复v4：桌面端恢复，移动端隐藏（第 5 轮）', '绝望尝试'],
        ['3ff4d90 (tag: v2)', 'absolute 包裹 + 水平滑动 + 圆点导航', '✅ 最终正确方案'],
    ]
)

add_heading('核心问题：AI 图片拉伸（5 轮才解决）', 3)
add_para('根因追踪：', bold=True)
for i, s in enumerate([
    '图片 1200×400，宽高比 3:1，容器在 flex 上下文中',
    'Flexbox 的 min-height: auto 允许元素按内容自然尺寸撑大',
    '图片 intrinsic size > 容器限制 → 溢出',
    'scrollIntoView 触发页面级滚动 → flexbox 重算 → 图片再拉伸',
    '全局 img 规则覆盖了 object-fit: cover',
], 1):
    add_numbered(s)

add_para('最终方案：position: absolute + aspect-ratio: 3/1。图片完全脱离 flex 文档流，容器尺寸由 aspect-ratio 严格约束，不受任何滚动重算影响。')
add_para('根本原因：AI 在跟 flexbox 的默认行为"对抗"而不是"利用"——应该在修复前打开 DevTools 量 computed height，而不是靠猜。')

add_table(
    ['轮次', '尝试方案', '失败原因'],
    [
        ['1', '全局 CSS 规则', '破坏已有布局'],
        ['2', 'flex 容器内修改', '没理解 min-height: auto'],
        ['3', '加 min-height: 0', '效果不一致'],
        ['4', '换 aspect-ratio', '没脱离文档流'],
        ['5', 'position: absolute', '✅ 有效'],
    ]
)

add_heading('导航栏 hamburger（3 轮失败）', 3)
add_table(
    ['轮次', '尝试', '反馈'],
    [
        ['1', '添加 hamburger + CSS 动画', '"白色横杠 + 黑线"'],
        ['2', '设透明背景', '"变成 3 条了"'],
        ['3', '回退到 v1（缩小字号+间距）', '✅ 满意'],
    ]
)
add_para('教训：毛玻璃 + neumorphism 深色阴影不适合做折叠菜单的视觉风格。6 项导航在 320px 下仍可见——折叠不是唯一选择。')

add_table(
    ['问题', '根因', '解决'],
    [
        ['圆点出现在错误位置', 'Python 替换占位符放错', '手动核对 HTML 嵌套'],
        ['圆点被拉长', 'button { min-height: 44px } 全局规则', '加 min-height: 0'],
        ['按钮触发图片拉伸', 'scrollIntoView 副作用', '改用 scrollTo + offsetLeft'],
        ['箭头按钮也触发拉伸', '同上', '移除箭头，手势 scroll-snap'],
        ['测试和正式文件不同步', '改了 test 文件没 merge', '直接改 index.html'],
    ]
)

add_heading('Phase 4：简历系统 · 2026-06-11（7 commits）', 2)
add_table(
    ['Commit', '内容'],
    [
        ['1de378e', '新增简历v2(英文Gmail+个人链接) + 邮箱跟随语言切换'],
        ['6aea643', '中文简历+网站：个人网站链接 + 英文简历v2上线'],
        ['228f0bb', '英文简历v2 PDF导出 + 下载链接改为PDF'],
        ['8438859', '英文简历v2替代旧版'],
        ['4b1efa9', '重新导出PDF：中文含个人网站链接、英文v2同步'],
        ['b27e9d3', '清理docs：删除旧版英文简历 + 废弃 full-cv 方案'],
    ]
)

add_para('关键沟通错误：', bold=True)
p = doc.add_paragraph()
run = p.add_run('错误 1：新增按钮 vs 替代旧版'); run.bold = True
p.add_run(' ——用户问"难道不应该替代之前的一页简历吗？你新增一个按钮是什么意思？"——没确认就动手。')

p = doc.add_paragraph()
run = p.add_run('错误 2：邮箱替换错了'); run.bold = True
p.add_run(' ——第一次把中文简历的 Foxmail 也改成了 Gmail。用户纠正：英文用 Gmail（国际送达保障），中文保留 Foxmail（国内直达）。最终用 JS 语言切换动态换 mailto: 和显示文本。')

add_para('docs 清理：删除 resume-en.html/pdf、full-cv-en-vA.html、full-cv-en-vC.html。保留 8 个文件：CAME 复盘（中/英）、完整 CV、简历（中/英 HTML+PDF）、工作职责导出、导出脚本、推荐信。')

# ═══════════════════════════════════════════════════════════
#  二、用户补充反馈分析（5 大块）
# ═══════════════════════════════════════════════════════════
add_heading('二、用户补充反馈深度分析', 1)
add_para('以下内容是用户在复盘基础上补充的 5 大块关键反馈，系统分析并附带调研结果和解决方案。')

# ── 板块 1 ──
add_heading('板块 1：版本管理与编号方案', 2)
add_para('问题描述：', bold=True)
add_para('版本回退有时退过了头，AI 无法准确识别时间线。每次改动应该用数字命名版本，让回退和修改更方便辨识。')

add_para('分析：', bold=True)
add_para('当前只有 2 个 tag（v1-desktop / v2-mobile-responsive），中间 10+ commit 没有可读版本号。比如 "AI 图片修了 5 轮"——应该标记为 v2.1、v2.2、v2.3……而不是靠 commit message 去翻。')

add_para('方案：版本编号规范（所有项目通用）', bold=True)
add_table(
    ['阶段', '编号', '示例'],
    [
        ['桌面版基线', 'v1', 'v1-desktop'],
        ['移动端响应式', 'v2', 'v2-mobile-responsive'],
        ['修复迭代', 'v2.x', 'v2.1（首次AI修）、v2.2（第二次AI修）'],
        ['功能新增', 'v3', 'v3-resume-system'],
        ['Bugfix', 'v3.x', 'v3.1-fix-email-switch'],
    ]
)
add_para('每个版本都在 git 打 tag。回退时指定 tag 名即可，不再需要翻 commit history。')

# ── 板块 2 ──
add_heading('板块 2：设计与 UI 工具链解决方案', 2)
add_para('问题描述：', bold=True)
add_para(
    'DeepSeek 出不了有审美的配色方案，AI 的 UI 设计也不令人满意。'
    '最终的成功是因为用户在网上下了一个固定提示词模板发给 AI，按那个模板做出来的——这不是 AI 独立的设计能力。'
    '需要调研其他人是怎么搭配工具链的（软件/插件/模型嵌套）。'
)

add_para('初版复盘遗漏的事实：', bold=True)
add_para('最终 Neumorphism 蓝方案的设计不是 AI 独立完成的。用户在网上找到了一个固定的提示词模板，提供给了 AI，AI 按照那个模板执行才得以成功。这说明 AI 在没有高质量设计参考输入时，无法独立产出有审美的方案。', italic=True)

add_para('调研结果：当前（2025-2026）最佳设计协作工具链', bold=True)

add_table(
    ['工具/方案', '核心能力', '适合场景'],
    [
        ['Claude Design', '上传 Figma/截图 → 自动抽取设计系统 → 多屏原型 → Handoff Bundle 交代码', '设计全流程'],
        ['Figma MCP', 'Claude 直接读 Figma 图层/颜色/间距/组件 → 代码和设计双向同步', '设计师 + AI 协作'],
        ['awesome-design-md (GitHub)', '58+ 品牌设计系统写成 DESIGN.md（Linear/Stripe/Apple 色值/字体/间距）', '设计参考 + prompt 输入'],
        ['shadcn/ui Figma Kit', 'Figma 组件库和代码 1:1 匹配', '前后端一致性'],
        ['DESIGN.md 放项目根目录', 'AI 直接读取颜色/字体/间距规范，一次定义永久遵守', '长期维护'],
    ]
)

add_para('核心结论：', bold=True)
add_para('没有 AI 能独立解决配色审美问题。正确的路径是：人类找参考（Dribbble / 品牌规范 / awesome-design-md）→ 把参考的结构化数据（DESIGN.md）喂给 AI → AI 在既定框架内执行。这要求团队在项目启动时就定义 Design System，而不是边做边想。')

# ── 板块 3 ──
add_heading('板块 3：过时信息导致的决策错误', 2)
add_para('问题描述：', bold=True)
add_para('Gitee Pages 服务器已经关闭，但 AI 没有查证就直接基于"Gitee Pages 可用"的假设进行工作，导致做了无用功。')

add_para('查证结果：', bold=True)
add_table(
    ['时间', '状态'],
    [
        ['2023-2024', 'Gitee Pages 开始整顿/限制'],
        ['2024-2025', '暂停服务'],
        ['2025', '确认停运，不再提供服务'],
    ]
)
add_para('根源：AI 的知识截止日期早于服务变更时间，但没有在行动前主动验证。')
add_para('方案：涉及第三方服务状态时，AI 必须先搜索确认当前状态再行动。用户也应该意识到 AI 的知识有时效性，涉及"这个服务还能不能用"的问题应直接提醒 AI 去查。')

# ── 板块 4 ──
add_heading('板块 4：技术开发循环绕圈与代码审查方案', 2)
add_para('问题描述：', bold=True)
add_para(
    '每次都是用户观察到根因、给出具体指令后 AI 才会去研究去做。反映三个问题：'
    '(1) 代码审核需要工具介入；(2) 问题排查方式需要系统化；(3) AI 绕圈圈（AI 图片修了 5 轮）——其他人怎么解决的？'
)

add_para('调研结果：解决 AI 绕圈的专用工具', bold=True)
add_table(
    ['工具', '原理', '适用'],
    [
        ['LoopLens MCP', '检测 AI 是否重复操作相同文件/命令 → 检测到循环就建议换方向', 'Claude Code 直接装 MCP'],
        ['Autonomous QA Loop', '每次审核用全新 agent（无历史记录）避免偏见', 'Prompt 模式'],
        ['No-No Debug', '自动记录每次错误/修复到 error_log.md，跨 session 累积规则', 'Skill 模式'],
        ['Pi Review Loop', '让 agent 自审到真的没问题才停，带 fresh context 模式', 'npm 包'],
        ['3-Prompt 管道', '意图提取 → Bug 寻找 → 修复建议，三步分开', '纯 prompt 技巧'],
    ]
)

add_para('代码审查工具：', bold=True)
add_table(
    ['工具', '做什么'],
    [
        ['antislope-ai', '每次保存自动本地审查（命名规范/边界风险/文档缺失）'],
        ['Plex', '在新上下文中审查代码，不继承作者偏见 → 发现真实 bug'],
        ['code-review-graph', '代码关系图谱，改一处自动分析波及范围'],
    ]
)

add_para('AI 绕圈的根因分析：', bold=True)
add_para(
    '当 AI 在一个问题上反复尝试但无效时，核心问题不是"不够努力"，而是判断方向的能力缺失。'
    '在 AI 图片修复案例中，每一轮都在改不同的 CSS 属性（flex / aspect-ratio / min-height），但没有一轮先打开 DevTools 确认当前各元素的实际 computed 值。'
    '解决方案：AI 在遇到 Bug 时，必须先"量化"（量尺寸、截图、读 computed style），再"定位"（找根因），最后"修复"（改代码）——这个顺序不能乱。'
)

# ── 板块 5 ──
add_heading('板块 5：文档读取编码问题（PDF/Word/Excel/PPTX）', 2)
add_para('问题描述：', bold=True)
add_para('AI 读取 PDF 时出现中文乱码，但有时选择绕过乱码、忽略重要信息直接出结果，而不是告诉用户哪里乱码或自行解决。')

add_para('当前 AI 对四种格式的读取能力：', bold=True)
add_table(
    ['格式', '中文读取', '英文读取', '问题'],
    [
        ['PDF', '⚠️ 有乱码', '✅ 正常', 'CID 编码 + 缺少 ToUnicode CMap → 解析库无法映射'],
        ['Word (.docx)', '✅ 正常', '✅ 正常', '—'],
        ['Excel (.xlsx)', '✅ 正常', '✅ 正常', '—'],
        ['PPTX', '✅ 正常', '✅ 正常', '—'],
    ]
)

add_para('PDF 乱码根因：', bold=True)
for t in [
    'PDF 使用 CID 编码（复杂字体）+ 缺少 ToUnicode CMap → 解析库无法将字形映射到 Unicode',
    '非嵌入字体：PDF 引用系统字体名但系统中没有 → 映射失败',
    '扫描件（图片型 PDF）完全没有文本层 → 任何文本提取工具均无效',
]:
    add_bullet(t)

add_para('修复方案（按优先级）：', bold=True)
for i, t in enumerate([
    '优先用 PyMuPDF 尝试提取，比 pdfplumber 对中文支持更好',
    '如果失败 → 用 PaddleOCR（深度学习 OCR，支持中英文混排）',
    '如果还是乱码 → 明确告知用户"此 PDF 第 X 页乱码"，而不是跳过',
], 1):
    add_numbered(t)

# ═══════════════════════════════════════════════════════════
#  三、PDF 排版留白问题专项分析
# ═══════════════════════════════════════════════════════════
add_heading('三、PDF 排版留白问题专项分析', 1)
add_para('用户反馈：简历和详细履历的 PDF 经常出现部分板块留白太多、字体大小不一致、内容有时全部挤压在一起有时又太过分散的问题。一页简历调了很久，详细履历因为页面太多暂时放弃。')

add_para('根因分析：', bold=True)

add_para('1. HTML → PDF 技术方案本身缺乏布局约束机制', bold=True)
add_para(
    '当前方案是用 HTML + CSS 写好页面结构，然后通过 Playwright 的 page.pdf() 将浏览器渲染结果输出为 PDF。'
    '浏览器对 flex/grid 的计算逻辑与 InDesign/Word 的排版逻辑完全不同：'
)
for t in [
    'Resume 的间距使用 2/4/6/8/10/12px 混用——不是任何倍数的关系',
    'Full CV 的 flex:1 意味着内容多就压缩、内容少就拉伸——留白多少是 flexbox 算出来的，不是设计决定的',
    '两个页面在 A4 上的内容高度不可预测，改一行字就可能让整个页面的 spacing 重新分布',
]:
    add_bullet(t)

add_para('2. 字体没有模数比例（Typographic Scale）', bold=True)
add_para(
    '英文简历里用了 22pt / 10.5pt / 9.5pt / 8.5pt / 8pt / 7.5pt 共 6 个层级，'
    '但这不是任何标准比例（如 1.25x Major Third）。视觉上"哪里不对"但说不出——因为没有层级间的比例关系。'
)

add_para('3. 多页 CV 的 flex 结构随内容变化不可控', bold=True)
add_para(
    '封面页内容少 → flex:1 拉伸 → 大片留白。ALTEN 页内容多 → flex 把所有东西挤在一起。'
    '没有 page-break-inside: avoid 控制分页行为，没有 min-height / max-height 约束。'
    '这就是为什么一页简历勉强能调（内容固定），但多页 CV 无法调好——每次修改的效果都依赖于当前内容量和浏览器计算。'
)

add_para('解决方案（双方案并行）：', bold=True)

p = doc.add_paragraph()
run = p.add_run('方案 A（首选）：Python → Word → PDF')
run.bold = True
add_para(
    'python-docx 可以精确控制：字体层级（pt 固定单位）、间距网格（cm/mm 绝对单位）、'
    '分页行为（page-break 精确指定）、页眉页脚。输出是"规范文档"级别，不是杂志排版级别，'
    '但 CV 使用完全足够。不会出现 flex 乱算导致的留白问题，改内容不会导致整个布局重新分布。'
)
add_para('技术验证：已成功用 python-docx 导出了 work-experience.docx（4 家公司工作职责，中英文混合，排版稳定）。')

p = doc.add_paragraph()
run = p.add_run('方案 B（备选）：Canva 排版')
run.bold = True
add_para(
    'AI 不能直接操作 Canva 界面，但可以把内容按 Canva 格式整理好（标题/正文/列表逐段标清），'
    '给出完整的排版规范（字体、字号、颜色、间距、每页结构），用户在 Canva 里拖拽粘贴即可。'
)

add_table(
    ['维度', '方案 A: Word→PDF', '方案 B: Canva'],
    [
        ['排版精度', '高（绝对单位控制）', '最高（可视化拖拽）'],
        ['AI 参与度', '全程自动化', '只能出规范和内容模板'],
        ['修改成本', '改代码重新运行即可', '手动调整'],
        ['适合场景', '简历/CV/正式文档', '作品集/展示型文档'],
        ['交付周期', '分钟级', '小时级（需人工操作）'],
    ]
)

# ═══════════════════════════════════════════════════════════
#  四、网格系统 + 字体模数比例 · 提前规划方案
# ═══════════════════════════════════════════════════════════
add_heading('四、网格系统 + 字体模数比例：提前规划方案', 1)
add_para('用户的建议：未来项目在设计阶段就提前把页面按网格划分，给到的版本有网格线条方便设计和 UI，上线前再隐藏网格。')
add_para('结论：完全可行，且是行业最佳实践。')

add_heading('网格系统（Grid Overlay）', 2)
add_para('设计阶段在页面上叠加半透明网格线，上线前隐藏。纯 CSS 实现：')

add_para('实现方式：', bold=True)
for t in [
    'CSS 中预设 .grid-overlay 类，默认 display: none',
    '设计阶段给 <html> 加一个 debug 类名（或 URL 参数 ?debug=true），网格显示',
    '上线前去掉即可——不需要删除代码，切换一个类名的事',
    '也可以用独立 grid-debug.html 做半透明叠加层，所有内容模块对齐到网格线',
]:
    add_bullet(t)

add_heading('字体模数比例（Typographic Scale）', 2)
add_para('定义一套有数学关系的字号层级，而不是凭感觉选大小。')
add_para('推荐比例：', bold=True)

add_table(
    ['比例名', '比值', '适合场景'],
    [
        ['Minor Second', '1.067', '连续文本密集型（如简历）'],
        ['Major Third', '1.250', '信息页面（标准网站）——推荐'],
        ['Perfect Fourth', '1.333', '展示型页面（作品集/营销页）'],
    ]
)

add_para('以 Major Third (1.25) 为例：', bold=True)
add_table(
    ['变量名', '计算', '大小', '用途'],
    [
        ['--text-xs', '16/1.25²', '~10px', '辅助标注'],
        ['--text-sm', '16/1.25', '~13px', '正文次要信息'],
        ['--text-base', '1rem', '16px', '正文'],
        ['--text-md', '16×1.25', '20px', '小标题'],
        ['--text-lg', '16×1.25²', '25px', '标题 H3'],
        ['--text-xl', '16×1.25³', '31px', '标题 H2'],
        ['--text-2xl', '16×1.25⁴', '39px', '标题 H1'],
        ['--text-3xl', '16×1.25⁵', '49px', 'Hero 大标题'],
    ]
)

add_heading('推荐的未来项目启动流程', 2)
add_table(
    ['步骤', '做什么', '产出'],
    [
        ['Step 1', '定网格：几列？间距多少？边距多少？', '12 列 + 20px gutter（推荐默认值）'],
        ['Step 2', '定字体比例：选一个模数', '1.25 (Major Third) 或 1.333 (Perfect Fourth)'],
        ['Step 3', '定义 Design Tokens：颜色/间距/圆角', 'CSS 变量表 + DESIGN.md'],
        ['Step 4', '你确认一次网格和比例', '签收后才能进入开发'],
        ['Step 5', '开发阶段：网格可见，模块对齐到网格', '.grid-overlay 显示'],
        ['Step 6', '上线前：隐藏网格', '.grid-overlay 隐藏或删除'],
    ]
)

add_para(
    '这套骨架搭好之后，用户只需要确认一次。后续所有修改都基于这套规则——不会再出现'
    'AI 问"缩小到多少"，用户说"1/3 还是 1/2"这种需要临时决策的情况。'
    '所有间距、字号、边距都是预先定好的，直接套用。'
)

# ═══════════════════════════════════════════════════════════
#  五、数据统计
# ═══════════════════════════════════════════════════════════
add_heading('五、数据统计', 1)
add_table(
    ['指标', '数值'],
    [
        ['总 commit', '16'],
        ['Git tags', '2（v1-desktop, v2-mobile-responsive）'],
        ['总代码行', '1633 行（纯 HTML+CSS+JS 单文件）'],
        ['模块数', '9'],
        ['断点层级', '4（900 / 768 / 600 / 480px）'],
        ['图片总数', '10 张（全部 loading="lazy"）'],
        ['简历版本', '2 套 4 文件（中/英 HTML + PDF）'],
        ['设计探索轮次', '8（最终靠用户提供的 prompt 模板成功）'],
        ['AI 图片修复轮次', '5（根因：未先量化再修）'],
        ['Hamburger 尝试轮次', '3（最终放弃，回退 v1 方案）'],
        ['部署目标', '2（GitHub Pages + Cloudflare Pages，Gitee Pages 已停运）'],
    ]
)

# ═══════════════════════════════════════════════════════════
#  六、深度问题分析
# ═══════════════════════════════════════════════════════════
add_heading('六、深度问题分析', 1)

add_heading('问题 1：没有阶段性验收就把 v1 推到了"完成"', 2)
add_para('v1 看起来"做完了"但移动端从未被测试。在打 v1-desktop tag 之前就应该声明"这是桌面版 v1，移动端尚未适配"。')

add_heading('问题 2：修复过程没有诊断数据', 2)
add_para('AI 图片修了 5 轮，每一轮都在"猜"——没有一次打开 DevTools 检查 computed height。正确的做法：先量化再修。')

add_heading('问题 3：设计探索和开发混在一起', 2)
add_para('8 套配色方案投入约 40% 时间但产出为零。核心矛盾：用户要参考和对比，但 AI 只能做成 HTML 才能展示。应该用截图/色板卡片/Figma 而非完整 HTML。')

add_heading('问题 4：沟通中的"做"和"想"没有分开', 2)
add_para('用户说"做一个看看"→ AI 理解为需求，做了完整方案。用户说"试试看"→ AI 花了 2 小时做出来。应该先确认方向再快速原型（<30min），而不是直接出完整成品。')

add_heading('问题 5：版本没有系统化管理', 2)
add_para('13 个修复 commit 只有 2 个 tag。回退只能靠翻 commit message。未来每个阶段性改动打数字 tag。')

add_heading('问题 6：HTML→PDF 方案不适合多页版式文档', 2)
add_para('浏览器 flex 计算与专业排版软件逻辑不同。一页简历勉强可调，多页 CV 无法稳定输出。首选 Python→Word→PDF，备选 Canva。')

add_heading('问题 7：第三方服务状态未经验证', 2)
add_para('Gitee Pages 已停运但 AI 没有主动查证就基于可用假设开展工作。涉及服务状态的问题必须先搜索确认。')

# ═══════════════════════════════════════════════════════════
#  七、做得好的地方
# ═══════════════════════════════════════════════════════════
add_heading('七、做得好的地方', 1)

items = [
    ('技术选型正确', [
        '纯 HTML/CSS/JS 单文件，部署极简，无构建工具，加载飞快',
        'data-i18n 双语方案优雅：天然 SEO 友好，切换即时无刷新',
    ]),
    ('版本控制意识', [
        'v1 打 tag 后才改 v2，用户可以随时 git checkout 回退',
        '每个修复阶段都有 commit，方便逐段回滚',
    ]),
    ('邮箱策略合理', [
        '英文 → Gmail（国际送达保障），中文 → Foxmail（国内直达）',
        '语言切换时动态换 mailto: 和显示文本，一个源文件维护',
    ]),
    ('最终 AI 图片方案经得起推敲', [
        'position: absolute + aspect-ratio + object-fit: cover',
        '所有浏览器行为一致，不受 flex 上下文影响，无需 JavaScript',
    ]),
]

for title, bullets in items:
    add_heading(title, 2)
    for b in bullets:
        add_bullet(b)

# ═══════════════════════════════════════════════════════════
#  八、SOP 提炼
# ═══════════════════════════════════════════════════════════
add_heading('八、SOP 提炼（下次直接套用）', 1)

add_heading('项目启动阶段', 2)
add_table(
    ['步骤', '做什么', '产出'],
    [
        ['1. 内容定稿', '收集文案/图片/数据，确认模块结构', '内容方案文档，签收后才能进入下一步'],
        ['2. 设计对齐', '出 2 个方向（截图/参考链接），用户选一个', 'DESIGN.md + 网格 + 字体比例'],
        ['3. 技术实现', '先桌面端 → 打 tag → 再移动端 → 逐一 breakpoint 测试', '每个阶段打数字 tag'],
        ['4. QA', '桌面 1920/1440/1024 + 平板 900/768 + 手机 600/480/375/320', '跨浏览器验证'],
    ]
)

add_heading('Bug 修复 SOP', 2)
for i, t in enumerate([
    '量化问题：量尺寸、截图、记录（打开 DevTools 看 computed values）',
    '定位根因：不要猜，用数据找出是 CSS 规则问题还是内容问题',
    '出一个最小复现：isolate the pattern',
    '提出方案 + 预期效果',
    '改一个地方，测一个地方',
    '确认不影响其他模块',
], 1):
    add_numbered(t)

add_heading('沟通规范', 2)
add_table(
    ['场景', '正确做法'],
    [
        ['"做一个看看"', '先确认方向，再快速原型（<30min），不是出完整方案'],
        ['"帮我改一下"', '重复确认需求："你的意思是……？"'],
        ['"试试这个方案"', '先说预期效果，再动手，动手前确认'],
        ['替代 vs 新增', '先问清楚再操作'],
        ['第三方服务状态', '先搜索确认当前状态，不依赖训练数据'],
    ]
)

add_heading('设计阶段网格 + 字体规范', 2)
add_table(
    ['项目', '默认推荐', '可选项'],
    [
        ['列数', '12 列', '8 列 / 6 列'],
        ['Gutter', '20px', '16px / 24px'],
        ['字体比例', 'Major Third (1.25)', 'Perfect Fourth (1.333) / Minor Second (1.067)'],
        ['基准字号', '16px (1rem)', '14px / 18px'],
        ['网格显示', '.grid-overlay + html.debug', '独立 grid-debug.html'],
    ]
)

# ═══════════════════════════════════════════════════════════
#  九、未完成事项
# ═══════════════════════════════════════════════════════════
add_heading('九、未完成事项', 1)
add_table(
    ['优先级', '事项', '状态'],
    [
        ['🟡 中', 'Dark Explorer 配色用户确认', '已搁置'],
        ['🟡 中', 'Full CV 排版（待用 Word→PDF 方案解决）', '待重新启动'],
        ['🟡 中', 'PDF 读取中文乱码问题（PaddleOCR 方案待实施）', '待测试'],
        ['🟢 低', '安装 LoopLens MCP / No-No Debug 等调试工具', '待配置'],
        ['🟢 低', '微信推文 9 宫格确认', '等待用户'],
        ['🟢 低', '头像/照片上传到 hero 区域', '需用户提供照片'],
        ['🔵 远期', '自定义域名绑定', '未开始'],
        ['🔵 远期', 'Python 后端 CMS', '未开始'],
    ]
)

# ── 页脚 ──
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('—— EOF ——')
run.font.color.rgb = RGBColor(0x5A, 0x6A, 0x76)
run.italic = True

output_path = r"C:\Users\64107\Desktop\portfolio-retrospective.docx"
doc.save(output_path)
print(f"OK → {output_path}  ({round(os.path.getsize(output_path)/1024)}K)")
